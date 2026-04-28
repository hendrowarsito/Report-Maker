import json
import re
from io import BytesIO
from itertools import zip_longest
import zipfile

import anthropic
import pandas as pd
import streamlit as st
from docx import Document

st.set_page_config(page_title="SRR Kalibata Report Maker", layout="wide")
st.markdown("# SRR Kalibata Report Maker")


# =============================================================================
# Konversi angka ke terbilang Rupiah (murni Python, tanpa API)
# =============================================================================

_SATUAN = [
    "", "satu", "dua", "tiga", "empat", "lima",
    "enam", "tujuh", "delapan", "sembilan",
]
_BELASAN = [
    "sepuluh", "sebelas", "dua belas", "tiga belas", "empat belas",
    "lima belas", "enam belas", "tujuh belas", "delapan belas", "sembilan belas",
]


def _terbilang_bilangan(n: int) -> str:
    """Rekursif: ubah bilangan bulat positif ke kata Indonesia (huruf kecil)."""
    if n == 0:
        return ""

    parts: list[str] = []

    if n >= 1_000_000_000_000:
        t, n = divmod(n, 1_000_000_000_000)
        parts.append(_terbilang_bilangan(t) + " triliun")

    if n >= 1_000_000_000:
        m, n = divmod(n, 1_000_000_000)
        parts.append(_terbilang_bilangan(m) + " miliar")

    if n >= 1_000_000:
        j, n = divmod(n, 1_000_000)
        parts.append(_terbilang_bilangan(j) + " juta")

    if n >= 1_000:
        r, n = divmod(n, 1_000)
        parts.append("seribu" if r == 1 else _terbilang_bilangan(r) + " ribu")

    if n >= 100:
        h, n = divmod(n, 100)
        parts.append("seratus" if h == 1 else _SATUAN[h] + " ratus")

    if n >= 20:
        p, n = divmod(n, 10)
        parts.append(_SATUAN[p] + " puluh")
        if n:
            parts.append(_SATUAN[n])
    elif n >= 10:
        parts.append(_BELASAN[n - 10])
    elif n > 0:
        parts.append(_SATUAN[n])

    return " ".join(parts)


def angka_ke_terbilang(n: int) -> str:
    """
    Ubah bilangan bulat (nilai Rupiah) ke terbilang huruf kapital.
    Contoh: 200_000_000 → 'DUA RATUS JUTA RUPIAH'
    """
    if n == 0:
        return "NOL RUPIAH"
    prefix = "MINUS " if n < 0 else ""
    return (prefix + _terbilang_bilangan(abs(n))).upper() + " RUPIAH"


def format_rupiah(n: int) -> str:
    """
    Format bilangan bulat ke string Rupiah Indonesia.
    Contoh: 200_000_000 → 'Rp 200.000.000,00'
    """
    ribuan = f"{n:,}".replace(",", ".")
    return f"Rp {ribuan},00"


def _parse_ke_int(nilai) -> int | None:
    """
    Coba parsing berbagai format nilai ke bilangan bulat.
    Mengembalikan None jika:
      - bukan angka
      - mengandung huruf selain prefix "Rp" (misal: tanggal, nama, %)
      - hasil < 1

    Format yang didukung: int, float, "200000000", "Rp 200.000.000,00", "200.000.000,00"
    """
    if isinstance(nilai, (int, float)):
        result = int(round(float(nilai)))
        return result if result >= 1 else None

    s = str(nilai).strip()

    # Tolak string yang mengandung huruf di luar prefix Rp, atau simbol non-numerik
    # Contoh yang DITOLAK: "28 April 2026", "Unit A", "20%", "n/a", "1.5x"
    tanpa_rp = re.sub(r"(?i)^\s*rp\.?\s*", "", s).strip()
    if re.search(r"[a-zA-Z%]", tanpa_rp):
        return None

    # Hapus prefix Rp
    s = tanpa_rp
    # Format Indonesia: titik = pemisah ribuan, koma = desimal
    if re.search(r"\d\.\d{3}", s):
        s = s.replace(".", "").replace(",", ".")
    else:
        s = s.replace(",", ".")

    s = re.sub(r"[^\d.]", "", s)
    if not s:
        return None
    try:
        result = int(round(float(s)))
        return result if result >= 1 else None
    except ValueError:
        return None


def format_angka_indonesia(n: int) -> str:
    """
    Format angka ke gaya Indonesia tanpa prefix Rp.
    Contoh: 200_000_000 → '200.000.000,00'
    Digunakan ketika template sudah hardcode 'Rp ' sebelum placeholder.
    """
    return f"{n:,}".replace(",", ".") + ",00"


def generate_auto_terbilang(
    data_dict: dict,
    terbilang_map: dict[str, str] | None = None,
) -> dict[str, dict]:
    """
    Buat entri terbilang & format Rupiah otomatis untuk setiap nilai numerik.

    terbilang_map  — pemetaan eksplisit dari key sumber ke nama placeholder
                     terbilang yang diinginkan, berasal dari kolom ke-3 Excel.
                     Contoh: {"NP_nominal": "NP_kalimat_cap"}
                     Jika None atau key tidak ada di map, fallback ke
                     suffix _TERBILANG / _FORMAT.

    Kembalikan dict entri baru saja (untuk ditampilkan di UI):
      { "NP_kalimat_cap": {"sumber": "NP_nominal", "angka": 200000000,
                           "terbilang": "DUA RATUS JUTA RUPIAH",
                           "format_rp": "Rp 200.000.000,00"} }
    """
    terbilang_map = terbilang_map or {}
    tambahan: dict[str, dict] = {}

    for key, val in data_dict.items():
        if key.endswith(("_TERBILANG", "_FORMAT")):
            continue

        angka = _parse_ke_int(val)
        if angka is None:
            continue

        info = {
            "sumber":       key,
            "angka":        angka,
            "terbilang":    angka_ke_terbilang(angka),
            "format_rp":    format_rupiah(angka),
            "format_angka": format_angka_indonesia(angka),
        }

        # Tentukan key terbilang: eksplisit dari map atau fallback suffix
        tb_key  = terbilang_map.get(key, key + "_TERBILANG")
        fmt_key = key + "_FORMAT"

        if tb_key not in data_dict:
            tambahan[tb_key] = info
        if fmt_key not in data_dict:
            tambahan[fmt_key] = info

    return tambahan


# =============================================================================
# Paragraph-level placeholder replacement — berbasis XML langsung
# =============================================================================
#
# MENGAPA XML, BUKAN para.runs?
# para.runs hanya mengembalikan <w:r> anak-langsung dari <w:p>.
# Placeholder bisa berada di dalam elemen pembungkus yang TIDAK terlihat oleh
# para.runs, antara lain:
#   • <w:hyperlink>   — teks yang merupakan hyperlink
#   • <w:ins>         — teks hasil tracked-change (insertion)
#   • <w:sdt>         — content control (structured document tag)
# Dengan para._element.iter(W_T) kita mendapatkan SEMUA <w:t> di seluruh
# sub-pohon XML, sehingga tidak ada placeholder yang terlewat.
# Karena kita hanya mengubah .text pada elemen <w:t> (bukan <w:rPr>),
# semua formatting (bold/italic/warna/font) tetap terjaga otomatis.
#
_W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_T   = f"{{{_W_NS}}}t"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _set_t_text(t_elem, text: str):
    """Tulis teks ke elemen <w:t>, pasang xml:space="preserve" jika ada spasi di tepi."""
    t_elem.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t_elem.set(_XML_SPACE, "preserve")


def _replace_in_para_xml(para, placeholder: str, value: str) -> int:
    """
    Ganti semua kemunculan placeholder di paragraf secara langsung via XML.
    Menangani: run normal, hyperlink, tracked-change, content-control,
               dan placeholder yang terpecah di beberapa elemen <w:t>.
    Kembalikan jumlah penggantian yang dilakukan.
    """
    t_elems = list(para._element.iter(_W_T))
    if not t_elems:
        return 0

    count = 0

    # Babak 1: ganti dalam satu elemen <w:t> (kasus paling umum)
    for t in t_elems:
        txt = t.text or ""
        if placeholder in txt:
            count += txt.count(placeholder)
            _set_t_text(t, txt.replace(placeholder, value))

    # Babak 2: ganti placeholder yang terpecah di beberapa <w:t>
    # Loop sampai tidak ada lagi kemunculan (menangani beberapa kejadian sekaligus)
    for _ in range(len(t_elems)):           # batas atas iterasi yang aman
        texts    = [t.text or "" for t in t_elems]
        full     = "".join(texts)
        if placeholder not in full:
            break

        ph_start = full.index(placeholder)
        ph_end   = ph_start + len(placeholder)

        # Bangun posisi kumulatif setiap elemen
        cum, cum_pos = 0, []
        for txt in texts:
            cum_pos.append(cum)
            cum += len(txt)

        s_idx = e_idx = None
        for i, (cp, txt) in enumerate(zip(cum_pos, texts)):
            cp_end = cp + len(txt)
            if cp <= ph_start < cp_end:
                s_idx = i
            if cp < ph_end <= cp_end:
                e_idx = i

        # Tangani placeholder yang tepat berakhir di batas elemen
        if e_idx is None:
            for i, cp in enumerate(cum_pos):
                if cp + len(texts[i]) == ph_end:
                    e_idx = i
                    break

        if s_idx is None or e_idx is None or s_idx == e_idx:
            break   # tidak dapat diselesaikan, hentikan agar tidak infinite loop

        before = texts[s_idx][: ph_start - cum_pos[s_idx]]
        after  = texts[e_idx][ph_end   - cum_pos[e_idx] :]

        _set_t_text(t_elems[s_idx], before + value)
        _set_t_text(t_elems[e_idx], after)
        for i in range(s_idx + 1, e_idx):
            t_elems[i].text = ""

        count += 1

    return count


def replace_placeholders_in_paragraph(para, data: dict):
    for key, value in data.items():
        _replace_in_para_xml(para, f"{{{{{key}}}}}", str(value))
    return para


# =============================================================================
# Document-level replacement (body + tabel tersarang + header + footer)
# =============================================================================

def _process_paragraph_collection(paragraphs, data: dict):
    for para in paragraphs:
        replace_placeholders_in_paragraph(para, data)


def _process_table_collection(tables, data: dict):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _process_paragraph_collection(cell.paragraphs, data)
                _process_table_collection(cell.tables, data)


def replace_placeholders(doc: Document, data: dict) -> Document:
    _process_paragraph_collection(doc.paragraphs, data)
    _process_table_collection(doc.tables, data)

    for section in doc.sections:
        for hf in (
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ):
            if hf is not None:
                _process_paragraph_collection(hf.paragraphs, data)
                _process_table_collection(hf.tables, data)

    return doc


# =============================================================================
# Find & Replace langsung (tanpa placeholder, tanpa API)
# =============================================================================

def apply_find_replace(doc: Document, pairs: list[tuple[str, str]]) -> int:
    """
    Ganti kata secara langsung di seluruh dokumen (body, tabel, header, footer).
    Menggunakan _replace_in_para_xml yang sama dengan placeholder replacement
    sehingga hyperlink, tracked-change, dan split antar <w:t> ikut tertangani.
    Kembalikan jumlah total penggantian yang dilakukan.
    """
    total = 0
    substitutions = [(f, r) for f, r in pairs if f.strip()]

    def _replace_in_para(para):
        nonlocal total
        for find, replace in substitutions:
            total += _replace_in_para_xml(para, find, replace)

    def _replace_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_para(para)
                    _replace_in_tables(cell.tables)

    for para in doc.paragraphs:
        _replace_in_para(para)
    _replace_in_tables(doc.tables)

    for section in doc.sections:
        for hf in (
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ):
            if hf is not None:
                for para in hf.paragraphs:
                    _replace_in_para(para)
                _replace_in_tables(hf.tables)

    return total


def smart_replace_with_claude(
    doc_text: str, find: str, replace: str, client: anthropic.Anthropic
) -> str:
    """
    Minta Claude menentukan apakah setiap kemunculan kata perlu diganti
    berdasarkan konteks kalimatnya (penggantian sensitif-konteks).
    Kembalikan penjelasan keputusan Claude sebagai teks markdown.
    Penggantian aktual tetap dilakukan oleh apply_find_replace —
    fungsi ini hanya memberi analisis konteks untuk dikonfirmasi user.
    """
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        messages=[
            {
                "role": "user",
                "content": (
                    f'Dalam dokumen laporan penilaian aset/properti Indonesia berikut, '
                    f'temukan setiap kemunculan kata **"{find}"** dan analisis:\n'
                    f'- Apakah dalam konteks ini penggantian dengan **"{replace}"** tepat?\n'
                    f'- Jika ada kemunculan yang TIDAK boleh diganti, jelaskan alasannya.\n\n'
                    f'Dokumen (penggal):\n---\n{doc_text[:8000]}\n---\n\n'
                    f'Jawab dalam markdown. Mulai dengan rekomendasi: '
                    f'✅ Ganti semua / ⚠️ Ganti sebagian / ❌ Jangan ganti'
                ),
            }
        ],
    ) as stream:
        return stream.get_final_text()


# =============================================================================
# File processing helpers
# =============================================================================

def process_files(word_file, data_dict: dict, find_replace_pairs: list[tuple[str, str]] | None = None):
    doc = Document(word_file)
    updated_doc = replace_placeholders(doc, data_dict)
    if find_replace_pairs:
        apply_find_replace(updated_doc, find_replace_pairs)
    output = BytesIO()
    updated_doc.save(output)
    output.seek(0)
    return output, updated_doc


def extract_text_from_docx(docx_obj: Document) -> str:
    """
    Ekstrak semua teks: paragraf tubuh dokumen + isi sel tabel.
    Versi lama hanya mengambil doc.paragraphs sehingga tabel tidak ikut dipreview.
    """
    parts: list[str] = []

    for para in docx_obj.paragraphs:
        parts.append(para.text)

    def _collect_table(table):
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = " | ".join(
                    para.text for para in cell.paragraphs if para.text.strip()
                )
                if cell_text:
                    row_cells.append(cell_text)
                for nested in cell.tables:
                    _collect_table(nested)
            if row_cells:
                parts.append("  ".join(row_cells))

    for table in docx_obj.tables:
        _collect_table(table)

    return "\n".join(parts)


def highlight_placeholders(text: str) -> str:
    return re.sub(r"(\{\{.*?\}\})", r"[PLACEHOLDER:\1]", text)


def scan_remaining_placeholders(text: str) -> list[str]:
    """Kembalikan daftar placeholder {{...}} yang belum terganti."""
    return sorted(set(re.findall(r"\{\{[^}]+\}\}", text)))


# =============================================================================
# Konsistensi nominal angka vs terbilang (regex — tidak butuh API)
# =============================================================================

def check_value_consistency_table(text: str) -> tuple[pd.DataFrame, list[int], list[str]]:
    """
    Kembalikan (dataframe, list angka, list kata) untuk keperluan validasi AI.

    Catatan bug yang diperbaiki vs versi lama:
    - Pasangan angka-kata sebelumnya hanya cek keberadaan, bukan kecocokan nilai.
      Sekarang disimpan terpisah agar bisa divalidasi ke Claude.
    """
    number_pattern = r"Rp\s*([\d\.]+,\d{2})\s*\("
    verbal_block_pattern = r"\(\s*([A-Z\s]+RUPIAH)\s*\)"

    raw_numbers = re.findall(number_pattern, text)
    number_ints: list[int] = []
    for n in raw_numbers:
        cleaned = n.replace(".", "").replace(",", ".")
        try:
            number_ints.append(int(float(cleaned)))
        except ValueError:
            pass

    verbal_blocks = re.findall(verbal_block_pattern, text.upper())
    normalized_verbal = [vb.strip() for vb in verbal_blocks]

    rows = []
    for angka, kata in zip_longest(number_ints, normalized_verbal, fillvalue="(tidak ditemukan)"):
        rows.append(
            {
                "Angka (Rp)": (
                    f"Rp {angka:,}".replace(",", ".")
                    if isinstance(angka, int)
                    else angka
                ),
                "Penyebutan dalam Kata": kata,
                "Keterangan (Regex)": (
                    "⚠️ Tidak cocok"
                    if angka == "(tidak ditemukan)" or kata == "(tidak ditemukan)"
                    else "❓ Belum divalidasi AI"
                ),
            }
        )

    return pd.DataFrame(rows), number_ints, normalized_verbal


# =============================================================================
# Claude API — tiga fitur utama
# =============================================================================

def _get_client(api_key: str) -> anthropic.Anthropic:
    return anthropic.Anthropic(api_key=api_key)


def validate_terbilang_ai(
    angka: int, terbilang: str, client: anthropic.Anthropic
) -> dict:
    """
    Validasi satu pasang angka-terbilang via Claude Haiku (cepat & hemat).
    Return: {"valid": bool | None, "koreksi": str | None}
    """
    try:
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=120,
            messages=[
                {
                    "role": "user",
                    "content": (
                        f"Validasi: apakah terbilang ini sesuai angkanya?\n"
                        f"Angka: Rp {angka:,}\n"
                        f"Terbilang: {terbilang}\n\n"
                        "Jawab hanya JSON tanpa markdown:\n"
                        '{"valid": true, "koreksi": null}\n'
                        "atau jika salah:\n"
                        '{"valid": false, "koreksi": "TERBILANG YANG BENAR RUPIAH"}'
                    ),
                }
            ],
        )
        text = response.content[0].text.strip()
        text = re.sub(r"^```json\s*|\s*```$", "", text, flags=re.MULTILINE).strip()
        return json.loads(text)
    except Exception as exc:
        return {"valid": None, "koreksi": None, "error": str(exc)}


def review_document_ai(doc_text: str, client: anthropic.Anthropic) -> str:
    """
    Review QA lengkap dokumen yang sudah terisi via Claude Sonnet (streaming).
    Cek: placeholder belum terisi, inkonsistensi nilai/tanggal, bagian kosong.
    """
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        messages=[
            {
                "role": "user",
                "content": (
                    "Anda adalah auditor dokumen laporan penilaian aset/properti profesional Indonesia.\n\n"
                    "Periksa dokumen berikut:\n"
                    "1. **Placeholder belum terisi** — pola {{...}} atau [PLACEHOLDER:{{...}}]\n"
                    "2. **Inkonsistensi nilai** — nominal angka yang sama berbeda di bagian lain\n"
                    "3. **Inkonsistensi tanggal** — tanggal tidak konsisten atau tidak logis\n"
                    "4. **Bagian kosong** — paragraf tidak lengkap\n"
                    "5. **Anomali data** — nilai tidak wajar, nama tidak konsisten\n\n"
                    f"Dokumen:\n---\n{doc_text[:10000]}\n---\n\n"
                    "Laporan audit dalam markdown terstruktur. "
                    "Mulai dengan baris status: ✅ LULUS / ⚠️ PERLU REVIEW / ❌ KRITIS"
                ),
            }
        ],
    ) as stream:
        return stream.get_final_text()


def analyze_template_ai(template_text: str, client: anthropic.Anthropic) -> dict:
    """
    Analisis template Word: deteksi semua placeholder dan artinya.
    Return dict: {"placeholders": {"KEY": "deskripsi"}, "wajib": [...], "opsional": [...]}
    """
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        messages=[
            {
                "role": "user",
                "content": (
                    "Analisis template laporan penilaian ini. "
                    "Temukan semua placeholder {{NAMA_PLACEHOLDER}} dan jelaskan:\n"
                    "- Artinya dalam konteks laporan penilaian properti/aset Indonesia\n"
                    "- Apakah wajib atau opsional\n"
                    "- Format data yang diharapkan\n\n"
                    f"Template:\n---\n{template_text[:8000]}\n---\n\n"
                    "Jawab hanya JSON tanpa markdown:\n"
                    '{"placeholders": {"KEY": {"arti": "...", "format": "...", "wajib": true}}, '
                    '"ringkasan": "..."}'
                ),
            }
        ],
    )
    text = response.content[0].text.strip()
    text = re.sub(r"^```json\s*|\s*```$", "", text, flags=re.MULTILINE).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return {"placeholders": {}, "ringkasan": text}


# =============================================================================
# Streamlit UI
# =============================================================================

with st.sidebar:
    st.header("📂 Upload Dokumen")
    word_files = st.file_uploader(
        "Template Laporan (.docx)", type=["docx"], accept_multiple_files=True
    )
    excel_file = st.file_uploader("Data Lembar Kerja (.xlsx)", type=["xlsx"])

    st.divider()
    st.header("🤖 Claude AI")
    api_key = st.text_input(
        "API Key Anthropic",
        type="password",
        placeholder="sk-ant-...",
        help="Dapatkan API key di console.anthropic.com",
    )
    claude_ready = bool(api_key and api_key.startswith("sk-ant-"))
    if api_key and not claude_ready:
        st.warning("Format API key tidak valid.")
    elif claude_ready:
        st.success("API key siap.")

    st.divider()
    st.header("⚙️ Konfigurasi Excel")
    sheet_name = st.text_input("Nama Sheet", value="Laporan")
    skiprows = st.number_input("Lewati baris pertama (skiprows)", min_value=0, value=6, step=1)

    st.divider()
    st.header("🔄 Find & Replace")
    st.caption("Ganti kata langsung di dokumen — tidak perlu placeholder.")

    if "fr_pairs" not in st.session_state:
        st.session_state.fr_pairs = [("", "")]

    fr_pairs_ui: list[tuple[str, str]] = []
    for i, (f_val, r_val) in enumerate(st.session_state.fr_pairs):
        c1, c2 = st.columns(2)
        find_val = c1.text_input("Cari", value=f_val, key=f"fr_find_{i}", label_visibility="collapsed", placeholder="Cari…")
        repl_val = c2.text_input("Ganti", value=r_val, key=f"fr_repl_{i}", label_visibility="collapsed", placeholder="Ganti dengan…")
        fr_pairs_ui.append((find_val, repl_val))

    if st.button("＋ Tambah baris"):
        st.session_state.fr_pairs = fr_pairs_ui + [("", "")]
        st.rerun()
    else:
        st.session_state.fr_pairs = fr_pairs_ui

    active_pairs = [(f, r) for f, r in fr_pairs_ui if f.strip()]


# ---------------------------------------------------------------------------
# Baca Excel & bangun data_dict
# ---------------------------------------------------------------------------

if excel_file:
    try:
        df_excel = pd.read_excel(
            excel_file,
            header=None,
            skiprows=int(skiprows),
            sheet_name=sheet_name,
            engine="openpyxl",
        )
    except Exception as e:
        st.error(f"Gagal membaca Excel: {e}")
        st.stop()

    st.subheader("📊 Tabel Data pada Lembar Kerja (edit jika perlu)")
    df_editable = st.data_editor(df_excel, num_rows="dynamic", use_container_width=True)

    data_dict: dict = {}
    for _, row in df_editable.iterrows():
        key = row.iloc[0]
        value = row.iloc[1]
        if pd.notnull(key):
            data_dict[str(key)] = str(value) if pd.notnull(value) else ""

    # Kolom ketiga (opsional): nama placeholder terbilang yang diinginkan
    # Contoh baris Excel: NP_nominal | 200000000 | NP_kalimat_cap
    terbilang_map: dict[str, str] = {}
    if df_editable.shape[1] >= 3:
        for _, row in df_editable.iterrows():
            key = row.iloc[0]
            tb_target = row.iloc[2]
            if pd.notnull(key) and pd.notnull(tb_target) and str(tb_target).strip():
                terbilang_map[str(key)] = str(tb_target).strip()

    # ---------------------------------------------------------------------------
    # Auto-generate terbilang & format Rupiah untuk semua nilai numerik
    # ---------------------------------------------------------------------------
    auto_tb = generate_auto_terbilang(data_dict, terbilang_map)

    if auto_tb:
        # Kelompokkan per key sumber agar tidak duplikat
        seen_src: set[str] = set()
        rows_tb = []

        for tb_key, info in auto_tb.items():
            src = info["sumber"]

            # Perbarui nilai source key: angka mentah → format Indonesia tanpa "Rp"
            # Sehingga template "Rp {{NP_nominal}}" menghasilkan "Rp 200.000.000,00"
            data_dict[src] = info["format_angka"]

            # Isi key terbilang (nama bisa bebas, dari map atau suffix _TERBILANG)
            if tb_key.endswith("_FORMAT"):
                data_dict[tb_key] = info["format_rp"]
            else:
                data_dict[tb_key] = info["terbilang"]

            # Isi _FORMAT selalu tersedia
            fmt_key = src + "_FORMAT"
            if fmt_key not in data_dict:
                data_dict[fmt_key] = info["format_rp"]

            if src in seen_src:
                continue
            seen_src.add(src)

            rows_tb.append({
                "Placeholder":           f"{{{{{src}}}}}",
                "Nilai → Diformat":      info["format_angka"],
                f"{{{{{src}_FORMAT}}}}" : info["format_rp"],
                "Placeholder Terbilang": f"{{{{{tb_key}}}}}",
                "Terbilang":             info["terbilang"],
            })

        # Tampilkan ringkasan di UI
        with st.expander("🔢 Terbilang & Format Rupiah yang Dibuat Otomatis", expanded=True):
            st.caption(
                "Kolom ke-3 Excel (opsional) → nama placeholder terbilang kustom. "
                "Kosong = nama otomatis `KEY_TERBILANG`."
            )
            if rows_tb:
                st.dataframe(pd.DataFrame(rows_tb), use_container_width=True)

    # ---------------------------------------------------------------------------
    # Fitur AI: Analisis Struktur Template
    # ---------------------------------------------------------------------------
    if word_files and claude_ready:
        with st.expander("🔍 Analisis Placeholder Template dengan AI", expanded=False):
            if st.button("Jalankan Analisis Template"):
                with st.spinner("Claude sedang membaca template..."):
                    client = _get_client(api_key)
                    all_texts = []
                    for wf in word_files:
                        doc_tmp = Document(wf)
                        all_texts.append(extract_text_from_docx(doc_tmp))
                        wf.seek(0)
                    combined = "\n\n".join(all_texts)
                    result = analyze_template_ai(combined, client)

                if result.get("placeholders"):
                    rows_analysis = []
                    for ph_key, info in result["placeholders"].items():
                        rows_analysis.append(
                            {
                                "Placeholder": f"{{{{{ph_key}}}}}",
                                "Arti": info.get("arti", ""),
                                "Format": info.get("format", ""),
                                "Wajib": "✅ Ya" if info.get("wajib") else "➖ Opsional",
                                "Terisi": "✅" if ph_key in data_dict and data_dict[ph_key] else "❌",
                            }
                        )
                    st.dataframe(pd.DataFrame(rows_analysis), use_container_width=True)
                    if result.get("ringkasan"):
                        st.info(result["ringkasan"])
                else:
                    st.write(result.get("ringkasan", "Tidak ada placeholder ditemukan."))

    # ---------------------------------------------------------------------------
    # Preview & proses setiap file Word
    # ---------------------------------------------------------------------------
    if word_files:
        st.subheader("📄 Preview Laporan — cek sebelum dibuat filenya")
        processed_files = []

        for word_file in word_files:
            output_file, updated_doc = process_files(word_file, data_dict, active_pairs or None)
            raw_text = extract_text_from_docx(updated_doc)
            preview_text = highlight_placeholders(raw_text)
            consistency_df, number_ints, normalized_verbal = check_value_consistency_table(
                raw_text
            )
            remaining = scan_remaining_placeholders(raw_text)

            with st.expander(f"📁 {word_file.name}", expanded=True):

                # -- Peringatan placeholder belum terisi (tanpa AI) --
                if remaining:
                    st.error(
                        f"**{len(remaining)} placeholder belum terganti:** "
                        + ", ".join(f"`{p}`" for p in remaining)
                    )

                # -- Preview teks --
                st.text_area(
                    "Isi Dokumen Terisi (Placeholder ditandai):",
                    value=preview_text,
                    height=400,
                    key=word_file.name + "_preview",
                    disabled=True,
                    label_visibility="collapsed",
                )

                # -- Tabel konsistensi nominal --
                if not consistency_df.empty:
                    st.markdown("🔍 **Cek Konsistensi Nominal Angka vs Kata**")

                    # Tombol validasi AI (hanya tampil jika API key tersedia)
                    if claude_ready and number_ints:
                        if st.button(
                            "✅ Validasi Terbilang dengan AI",
                            key=word_file.name + "_val_btn",
                        ):
                            client = _get_client(api_key)
                            ai_results = []
                            progress = st.progress(0, text="Memvalidasi dengan Claude...")
                            for i, (angka, kata) in enumerate(
                                zip_longest(number_ints, normalized_verbal, fillvalue="(tidak ditemukan)")
                            ):
                                if isinstance(angka, int) and kata != "(tidak ditemukan)":
                                    res = validate_terbilang_ai(angka, kata, client)
                                    if res.get("valid") is True:
                                        ai_results.append("✅ Benar")
                                    elif res.get("valid") is False:
                                        koreksi = res.get("koreksi", "")
                                        ai_results.append(f"❌ Salah — Seharusnya: {koreksi}")
                                    else:
                                        ai_results.append("⚠️ Tidak dapat divalidasi")
                                else:
                                    ai_results.append("⚠️ Tidak cocok")
                                progress.progress(
                                    (i + 1) / max(len(number_ints), len(normalized_verbal), 1)
                                )
                            progress.empty()

                            consistency_df["Validasi AI"] = ai_results
                            st.session_state[word_file.name + "_consistency"] = consistency_df

                    # Tampilkan tabel (dari session state jika sudah divalidasi AI)
                    display_df = st.session_state.get(
                        word_file.name + "_consistency", consistency_df
                    )
                    st.dataframe(display_df, use_container_width=True)

                    has_error = "⚠️" in display_df.get("Keterangan (Regex)", pd.Series()).values or (
                        "Validasi AI" in display_df.columns
                        and display_df["Validasi AI"].str.contains("❌|⚠️").any()
                    )
                    if has_error:
                        st.warning("Terdapat ketidaksesuaian nominal — periksa sebelum diunduh.")
                    elif "Validasi AI" in display_df.columns:
                        st.success("Semua nominal tervalidasi AI — sesuai.")

                # -- Review dokumen lengkap dengan AI --
                if claude_ready:
                    st.markdown("---")
                    if st.button(
                        "🤖 Review Lengkap Dokumen dengan AI",
                        key=word_file.name + "_review_btn",
                    ):
                        client = _get_client(api_key)
                        with st.spinner("Claude sedang mereview dokumen..."):
                            review_result = review_document_ai(raw_text, client)
                        st.session_state[word_file.name + "_review"] = review_result

                    if word_file.name + "_review" in st.session_state:
                        st.markdown("**Hasil Review AI:**")
                        st.markdown(st.session_state[word_file.name + "_review"])

                # -- Analisis konteks Find & Replace dengan Claude --
                if claude_ready and active_pairs:
                    st.markdown("---")
                    st.markdown("**🔄 Analisis Konteks Find & Replace**")
                    for find_word, replace_word in active_pairs:
                        btn_key = f"{word_file.name}_fr_{find_word}"
                        if st.button(
                            f'Analisis: "{find_word}" → "{replace_word}"',
                            key=btn_key,
                        ):
                            client = _get_client(api_key)
                            with st.spinner(f'Claude menganalisis konteks "{find_word}"...'):
                                analysis = smart_replace_with_claude(
                                    raw_text, find_word, replace_word, client
                                )
                            st.session_state[btn_key + "_result"] = analysis

                        if btn_key + "_result" in st.session_state:
                            st.markdown(st.session_state[btn_key + "_result"])

            processed_files.append((word_file.name, output_file))

        # ---------------------------------------------------------------------------
        # Unduh hasil
        # ---------------------------------------------------------------------------
        st.subheader("⬇️ Unduh Hasil")
        if len(processed_files) == 1:
            file_name, file_obj = processed_files[0]
            st.download_button(
                label="Unduh Dokumen",
                data=file_obj,
                file_name=file_name.replace(".docx", "_terisi.docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                for name, file_obj in processed_files:
                    filename = name.replace(".docx", "_terisi.docx")
                    zipf.writestr(filename, file_obj.getvalue())
            zip_buffer.seek(0)
            st.download_button(
                label="Unduh Semua Dokumen (.zip)",
                data=zip_buffer,
                file_name="Semua_Dokumen_Terisi.zip",
                mime="application/zip",
            )
